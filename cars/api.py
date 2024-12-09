from rest_framework.viewsets import GenericViewSet
from rest_framework import mixins, viewsets

from rest_framework.decorators import action
from rest_framework.response import Response

from cars.models import Car, Mark, CarClass, BodyType, Country, UserPhoto
from cars.serializers import CarCreateSerializer, CarSerializer, MarkSerializer, CarClassSerializer, BodyTypeSerializer, CountrySerializer, UserPhotoSerializer
from django.db.models import Avg, Count, Max, Min
from rest_framework import serializers
from django.contrib.auth import authenticate, login

import openpyxl
from io import BytesIO
from docx import Document
from django.http import HttpResponse

class CarsViewset(
    mixins.CreateModelMixin, 
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin, 
    mixins.ListModelMixin, 
    mixins.DestroyModelMixin, 
    GenericViewSet
    ):

    queryset = Car.objects.all()
    serializer_class = CarSerializer

    def get_serializer_class(self):
        if self.action in ('update', 'create'):
            return CarCreateSerializer
        return super().get_serializer_class()
    
    def get_queryset(self):
        qs = super().get_queryset()

        if (self.request.user.is_superuser):
            resQs = qs
        else:
            # фильтруем по текущему юзеру
            resQs = qs.filter(user=self.request.user)

        return resQs
    
    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()
        most_mark_name = serializers.CharField()
        most_body_type = serializers.CharField()
        most_class = serializers.CharField()
        most_country = serializers.CharField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        most_common_mark = Car.objects.values('mark_name').annotate(mark_count=Count('id')).order_by('-mark_count').first()
        most_common_body_type = Car.objects.values('body_type').annotate(body_type_count=Count('id')).order_by('-body_type_count').first()
        most_common_class = Car.objects.values('car_class').annotate(car_class_count=Count('id')).order_by('-car_class_count').first()
        most_common_country = Car.objects.values('country').annotate(country_count=Count('id')).order_by('-country_count').first()
            
        stats = Car.objects.aggregate(
            count=Count("*"), 
            avg=Avg("id"), 
            max=Max("id"), 
            min=Min("id"), 
        )

        if most_common_mark:
            stats['most_mark_name'] = Mark.objects.get(pk=most_common_mark['mark_name'])
        else:
            stats['most_mark_name'] = None

        if most_common_body_type:
            stats['most_body_type'] = BodyType.objects.get(pk=most_common_body_type['body_type'])
        else:
            stats['most_body_type'] = None

        if most_common_class:
            stats['most_class'] = CarClass.objects.get(pk=most_common_class['car_class'])
        else:
            stats['most_class'] = None

        if most_common_country:
            stats['most_country'] = Country.objects.get(pk=most_common_country['country'])
        else:
            stats['most_country'] = None

        serializer = self.StatsSerializer(instance = stats)

        return Response(serializer.data)
    
    @action(detail=False, methods=["GET"], url_path="export-excel")
    def export_to_excel(self, request):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Cars"

        headers = ["ID", "Model", "Mark", "CarClass", "BodyType", "Country"]
        ws.append(headers)

        cars = Car.objects.all()

        user = request.user
        if user.is_superuser == False:
            cars = list(filter(lambda car: car.username == user.username, cars))


        for car in cars:
            ws.append([
                car.id,
                car.model,
                car.mark_name.name,
                car.car_class.name,
                car.body_type.name,
                car.country.name
            ])

        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = HttpResponse(output, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response["Content-Disposition"] = 'attachment; filename="cars.xlsx"'
        return response

    @action(detail=False, methods=["GET"], url_path="export-word")
    def export_to_word(self, request):
        doc = Document()
        doc.add_heading("Cars", level=1)

        table = doc.add_table(rows=1, cols=6)
        headers = ["ID", "Model", "Mark", "CarClass", "BodyType", "Country"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        cars = Car.objects.all()

        user = request.user
        if user.is_superuser == False:
            cars = list(filter(lambda car: car.username == user.username, cars))
            
        for car in cars:
            row_cells = table.add_row().cells
            row_cells[0].text = str(car.id)
            row_cells[1].text = car.model
            row_cells[2].text = car.mark_name.name
            row_cells[3].text = car.car_class.name
            row_cells[4].text = car.body_type.name
            row_cells[5].text = car.country.name

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        response = HttpResponse(output, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = 'attachment; filename="cars.docx"'
        return response

class MarksViewset(
    mixins.CreateModelMixin, 
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin, 
    mixins.ListModelMixin, 
    mixins.DestroyModelMixin, 
    GenericViewSet
    ):

    queryset = Mark.objects.all()
    serializer_class = MarkSerializer

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
            
        stats = Mark.objects.aggregate(
            count=Count("*"),
        )

        serializer = self.StatsSerializer(instance = stats)

        return Response(serializer.data)

class CarClassesViewset(
    mixins.CreateModelMixin, 
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin, 
    mixins.ListModelMixin, 
    mixins.DestroyModelMixin, 
    GenericViewSet
    ):

    queryset = CarClass.objects.all()
    serializer_class = CarClassSerializer

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
            
        stats = CarClass.objects.aggregate(
            count=Count("*"),
        )

        serializer = self.StatsSerializer(instance = stats)

        return Response(serializer.data)

class BodyTypesViewset(
    mixins.CreateModelMixin, 
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin, 
    mixins.ListModelMixin, 
    mixins.DestroyModelMixin, 
    GenericViewSet
    ):

    queryset = BodyType.objects.all()
    serializer_class = BodyTypeSerializer

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
            
        stats = BodyType.objects.aggregate(
            count=Count("*"),
        )

        serializer = self.StatsSerializer(instance = stats)

        return Response(serializer.data)

class CountriesViewset(
    mixins.CreateModelMixin, 
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin, 
    mixins.ListModelMixin, 
    mixins.DestroyModelMixin, 
    GenericViewSet
    ):

    queryset = Country.objects.all()
    serializer_class = CountrySerializer

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
            
        stats = Country.objects.aggregate(
            count=Count("*"),
        )

        serializer = self.StatsSerializer(instance = stats)

        return Response(serializer.data)

class UserProfileViewset(GenericViewSet):
    @action(url_path="info", detail=False, methods=["GET"])
    def get_url(self, request, *args, **kwargs):
        user = request.user
        data = {
            "is_authentificated": user.is_authenticated
        }
        if user.is_authenticated:
            data.update({
                "is_superuser": user.is_superuser, 
                "name": user.username
            })
        return Response(data)
    
    @action(url_path="login", detail=False, methods=["POST"])
    def login(self, request, *args, **kwargs):
        user = authenticate(username=self.request.data['username'], password=self.request.data['password'])
        success = False
        if user is not None:
            login(request, user)
            success  = True

        return Response({
            "success": success 
        })
    
class UsersPhotoViewset(
    mixins.CreateModelMixin, 
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin, 
    mixins.ListModelMixin, 
    mixins.DestroyModelMixin, 
    GenericViewSet
    ):

    queryset = UserPhoto.objects.all()
    serializer_class = UserPhotoSerializer